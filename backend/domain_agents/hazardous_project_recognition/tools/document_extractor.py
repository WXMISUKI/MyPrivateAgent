"""
文档内容提取工具

支持从 .doc/.docx/.xlsx/.xls/.csv 文件中提取文本和表格内容。
用于危大工程识别等需要读取文档的 Agent。
"""

from __future__ import annotations

import csv
import io
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_document_content(file_path: str) -> dict:
    """从文档文件中提取文本和表格内容。

    Args:
        file_path: 文件路径

    Returns:
        {
            "status": "success" | "error",
            "filename": "文件名",
            "file_type": "docx|xlsx|xls|csv|doc",
            "text_content": "全文文本",
            "tables": [
                {
                    "sheet_name": "Sheet1",
                    "headers": ["列1", "列2"],
                    "rows": [["值1", "值2"], ...]
                }
            ],
            "page_count": 2,
            "error": "错误信息（仅 status=error 时）"
        }
    """
    path = Path(file_path)
    if not path.exists():
        return {"status": "error", "error": f"文件不存在: {file_path}"}

    suffix = path.suffix.lower()
    filename = path.name

    try:
        if suffix in (".docx", ".doc"):
            return _extract_word(file_path, filename)
        elif suffix in (".xlsx", ".xls"):
            return _extract_excel(file_path, filename)
        elif suffix == ".csv":
            return _extract_csv(file_path, filename)
        else:
            return {"status": "error", "error": f"不支持的文件格式: {suffix}"}
    except Exception as e:
        logger.error(f"文档提取失败: {e}")
        return {"status": "error", "error": str(e), "filename": filename}


def _extract_word(file_path: str, filename: str) -> dict:
    """从 Word 文档提取内容。支持 .docx 和旧版 .doc 格式。"""
    path = Path(file_path)

    # 先尝试 python-docx（.docx 格式）
    try:
        from docx import Document
        doc = Document(file_path)

        # 提取全文
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 提取表格
        tables = []
        for i, table in enumerate(doc.tables):
            headers = []
            rows = []
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if j == 0:
                    headers = cells
                else:
                    rows.append(cells)
            if headers:
                tables.append({
                    "sheet_name": f"Table_{i+1}",
                    "headers": headers,
                    "rows": rows,
                })

        return {
            "status": "success",
            "filename": filename,
            "file_type": "docx",
            "text_content": "\n".join(text_parts),
            "tables": tables,
            "page_count": len(doc.sections) if doc.sections else 1,
        }
    except Exception as e:
        # python-docx 失败，尝试 .doc 格式（旧版二进制格式）
        logger.info(f"python-docx failed for {filename}, trying legacy .doc parsing: {e}")

    # 尝试用 olefile 或 textract 处理旧版 .doc
    try:
        return _extract_legacy_doc(file_path, filename)
    except Exception as e2:
        return {
            "status": "error",
            "error": f"Word 文档解析失败: python-docx 错误={e}, legacy 解析错误={e2}",
            "filename": filename,
        }


def _extract_legacy_doc(file_path: str, filename: str) -> dict:
    """从旧版 .doc 格式提取内容（使用 olefile 或直接文本提取）。"""
    import subprocess
    import tempfile

    # 方法 1: 尝试用 LibreOffice 转换
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", "--outdir", tempfile.gettempdir(), file_path],
            capture_output=True, timeout=30,
        )
        if result.returncode == 0 and Path(tmp_path).exists():
            from docx import Document
            doc = Document(tmp_path)
            text_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            Path(tmp_path).unlink(missing_ok=True)
            return {
                "status": "success",
                "filename": filename,
                "file_type": "doc",
                "text_content": "\n".join(text_parts),
                "tables": [],
                "page_count": 1,
            }
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 方法 2: 直接读取二进制提取可读文本
    try:
        with open(file_path, "rb") as f:
            raw = f.read()

        # 从 OLE 复合文档中提取 UTF-16LE 编码的文本
        text_parts = []
        # 简单的文本提取：查找连续的可打印字符
        current_text = []
        i = 0
        while i < len(raw) - 1:
            # UTF-16LE 字符
            char_code = raw[i] | (raw[i + 1] << 8)
            if 0x20 <= char_code <= 0x7E or 0x4E00 <= char_code <= 0x9FFF or char_code in (0x0A, 0x0D, 0x09):
                current_text.append(chr(char_code))
            else:
                if len(current_text) > 3:
                    text = "".join(current_text).strip()
                    if text and not text.isspace():
                        text_parts.append(text)
                current_text = []
            i += 2

        if current_text:
            text = "".join(current_text).strip()
            if text and len(text) > 3:
                text_parts.append(text)

        if text_parts:
            return {
                "status": "success",
                "filename": filename,
                "file_type": "doc",
                "text_content": "\n".join(text_parts),
                "tables": [],
                "page_count": 1,
            }
    except Exception:
        pass

    return {
        "status": "error",
        "error": "无法解析 .doc 格式。建议转换为 .docx 后重试。",
        "filename": filename,
    }


def _extract_excel(file_path: str, filename: str) -> dict:
    """从 Excel 文件提取内容。"""
    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=True)
    text_parts = []
    tables = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data = []
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in row_values):
                rows_data.append(row_values)

        if not rows_data:
            continue

        # 智能识别表头行：跳过只有第一个单元格有内容的行（标题行）
        header_idx = 0
        for i, row in enumerate(rows_data):
            non_empty = [v for v in row if v.strip()]
            if len(non_empty) > 1:
                header_idx = i
                break

        headers = rows_data[header_idx]
        data_rows = rows_data[header_idx + 1:]

        # 过滤空行
        data_rows = [r for r in data_rows if any(v.strip() for v in r)]

        if headers:
            tables.append({
                "sheet_name": sheet_name,
                "headers": headers,
                "rows": data_rows,
            })
            # 构建文本（带表头）
            text_parts.append(" | ".join(headers))
            for row in data_rows:
                text_parts.append(" | ".join(row))

    wb.close()

    return {
        "status": "success",
        "filename": filename,
        "file_type": "xlsx",
        "text_content": "\n".join(text_parts),
        "tables": tables,
        "page_count": len(tables),
    }


def _extract_csv(file_path: str, filename: str) -> dict:
    """从 CSV 文件提取内容。"""
    # 尝试多种编码
    content = None
    for encoding in ["utf-8", "gbk", "gb2312", "utf-8-sig"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        return {"status": "error", "error": "CSV 文件编码无法识别", "filename": filename}

    reader = csv.reader(io.StringIO(content))
    rows_data = [row for row in reader if any(cell.strip() for cell in row)]

    headers = rows_data[0] if rows_data else []
    data_rows = rows_data[1:] if len(rows_data) > 1 else []
    text_parts = [" | ".join(row) for row in rows_data]

    return {
        "status": "success",
        "filename": filename,
        "file_type": "csv",
        "text_content": "\n".join(text_parts),
        "tables": [{
            "sheet_name": "Sheet1",
            "headers": headers,
            "rows": data_rows,
        }],
        "page_count": 1,
    }


def validate_json_output(json_str: str) -> dict:
    """验证 JSON 输出是否符合危大工程清单格式。

    Args:
        json_str: 待验证的 JSON 字符串

    Returns:
        {"valid": true/false, "errors": [...], "data": [...]}
    """
    try:
        data = json.loads(json_str)
    except json.JSONDecodeError as e:
        return {"valid": False, "errors": [f"JSON 解析失败: {e}"], "data": None}

    errors = []

    if not isinstance(data, dict):
        return {"valid": False, "errors": ["输出必须是 JSON 对象"], "data": None}

    if "code" not in data:
        errors.append("缺少 code 字段")
    if "msg" not in data:
        errors.append("缺少 msg 字段")

    if data.get("code") == 200 and "data" in data:
        items = data["data"]
        if not isinstance(items, list):
            errors.append("data 必须是数组")
        else:
            for i, item in enumerate(items):
                if not isinstance(item, dict):
                    errors.append(f"data[{i}] 必须是对象")
                    continue
                for field in ["originname", "name", "category", "isExdanger"]:
                    if field not in item:
                        errors.append(f"data[{i}] 缺少字段 {field}")
                if "isExdanger" in item and not isinstance(item["isExdanger"], bool):
                    errors.append(f"data[{i}].isExdanger 必须是布尔值")

    return {
        "valid": len(errors) == 0,
        "errors": errors,
        "data": data.get("data") if data.get("code") == 200 else None,
    }
